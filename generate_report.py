"""
生成汽车理论课程设计报告 DOCX
封面/题目描述来自参考DOCX，目录在最前，每题代码附在题后
"""
import os, zipfile, re, io
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

doc = Document()

# ============ 页面 A4 ============
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ============ 字体工具 ============
def ensure_rFonts(run, ea='宋体', latin='宋体'):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), ea)
    rFonts.set(qn('w:ascii'), latin)
    rFonts.set(qn('w:hAnsi'), latin)
    for attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme']:
        k = qn(attr)
        if k in rFonts.attrib:
            del rFonts.attrib[k]

def set_run(run, font='宋体', size=None, bold=False):
    ensure_rFonts(run, font, font)
    if size: run.font.size = size
    if bold: run.bold = True

def add_para(text, font='宋体', size=Pt(12), bold=False, align=None):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    run = p.add_run(text)
    set_run(run, font, size, bold)
    return p

def add_empty_line():
    doc.add_paragraph()

def add_formula(formula_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(formula_text)
    ensure_rFonts(run, 'Cambria Math', 'Cambria Math')
    run.font.size = Pt(12)
    return p

def add_code_block(code_text):
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(14)
        run = p.add_run(line)
        ensure_rFonts(run, '宋体', 'Consolas')
        run.font.name = 'Consolas'
        run.font.size = Pt(8)

def add_result_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run(run, '宋体', Pt(10), True)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_run(run, '宋体', Pt(10))
    add_empty_line()
    return table

def fix_style_rFonts(style):
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme']:
        k = qn(attr)
        if k in rFonts.attrib:
            del rFonts.attrib[k]
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), '宋体')
    rFonts.set(qn('w:hAnsi'), '宋体')

def add_page_break():
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

# 修复样式
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
fix_style_rFonts(style)
for i in range(1, 4):
    h_style = doc.styles[f'Heading {i}']
    h_style.font.name = '宋体'
    h_style.font.color.rgb = RGBColor(0, 0, 0)
    fix_style_rFonts(h_style)
    h_style.font.size = {1: Pt(16), 2: Pt(14), 3: Pt(13)}[i]

def add_toc():
    """动态目录页"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('目  录')
    set_run(run, '宋体', Pt(18), True)
    add_empty_line()

    p_toc = doc.add_paragraph()
    for tag, attrs in [
        ('w:fldChar', {qn('w:fldCharType'): 'begin'}),
        ('w:instrText', {qn('xml:space'): 'preserve', 'text': ' TOC \\o "1-3" \\h \\z \\u '}),
        ('w:fldChar', {qn('w:fldCharType'): 'separate'}),
    ]:
        run = p_toc.add_run()
        el = OxmlElement(tag)
        for k, v in attrs.items():
            if k == 'text':
                el.text = v
            else:
                el.set(k, v)
        run._r.append(el)
    run_text = p_toc.add_run('（请按 Ctrl+A 后按 F9 更新目录）')
    set_run(run_text, '宋体', Pt(10))
    run_end = p_toc.add_run()
    el_end = OxmlElement('w:fldChar')
    el_end.set(qn('w:fldCharType'), 'end')
    run_end._r.append(el_end)
    add_page_break()

def add_problem_cover(number, title_cn):
    """每个题目的独立封面"""
    for _ in range(6):
        add_empty_line()
    add_para(f'汽车理论课程设计{number}', '宋体', Pt(22), True, WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(4):
        add_empty_line()
    add_para('专业名称：车辆工程', '宋体', Pt(14), False, WD_ALIGN_PARAGRAPH.CENTER)
    add_para('班    级：2023370501', '宋体', Pt(14), False, WD_ALIGN_PARAGRAPH.CENTER)
    add_para('学    号：2023901021', '宋体', Pt(14), False, WD_ALIGN_PARAGRAPH.CENTER)
    add_para('学生姓名：李思成', '宋体', Pt(14), False, WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line()
    add_para('2026 年 6 月 1 日', '宋体', Pt(14), False, WD_ALIGN_PARAGRAPH.CENTER)
    add_page_break()
    # 题目标题
    add_para(f'《汽车理论》课程设计{number}（{title_cn}）', '宋体', Pt(14), True, WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line()

# ============================================================
# 目录（最前）
# ============================================================
add_toc()

# ============================================================
# 题目一
# ============================================================
add_problem_cover('一', '习题1.3')
add_para('1.3 题目：确定一轻型货车的动力性能（货车可装用4挡或5挡变速器，任选其中的一种进行整车性能计算）：')
add_para('1）绘制汽车驱动力与行驶阻力平衡图。')
add_para('2）求汽车最高车速，最大爬坡度及克服该坡度时相应的附着率。')
add_para('3）绘制汽车行驶加速度倒数曲线，用图解积分法求汽车用2档起步加速行驶至70km/h的车速－时间曲线，或者用计算机求汽车用2档起步加速行驶至70km/h的加速时间。')
add_para('轻型货车的有关数据如下：')
add_para('汽油发动机使用外特性的Tq-n曲线的拟合公式为：')
add_formula('Tq = -19.313 + 295.27(n/1000) - 165.44(n/1000)² + 40.874(n/1000)³ - 3.8445(n/1000)⁴')
add_para('式中，Tq为发动机转矩（N·m）；n为发动机转速（r/min）。')
add_para('发动机的最低转速n_min=600r/min，最高转速n_max=4000r/min。')
add_para('变速器传动比ig（五档）：5.56 / 2.769 / 1.644 / 1.00 / 0.793')
add_para('（注：本题最好上机计算，可以选择Matlab、Python、Excel、VB、VC等软件进行计算和画图。）')

add_result_table(['参数', '数值'], [
    ['装载质量', '2000 kg'], ['整车整备质量', '1800 kg'], ['总质量', '3880 kg'],
    ['车轮半径', '0.367 m'], ['传动系机械效率 ηt', '0.85'],
    ['滚动阻力系数 f', '0.013'], ['空气阻力系数×迎风面积 CDA', '2.77 m²'],
    ['主减速器传动比 i0', '5.83'], ['飞轮转动惯量 If', '0.218 kg·m²'],
    ['二前轮转动惯量 Iw1', '1.798 kg·m²'], ['四后轮转动惯量 Iw2', '3.598 kg·m²'],
    ['轴距 L', '3.2 m'], ['质心至前轴距离（满载）a', '1.947 m'],
    ['质心高（满载）hg', '0.9 m'],
])

doc.add_heading('解题过程', level=1)
doc.add_heading('1. 计算公式', level=2)
for label, formula in [
    ('发动机外特性转矩：', 'Tq = -19.313 + 295.27(n/1000) - 165.44(n/1000)² + 40.874(n/1000)³ - 3.8445(n/1000)⁴'),
    ('各档车速：', 'ua = 0.377 × r × n / (ig × i0)  [km/h]'),
    ('各档驱动力：', 'Ft = Tq × ig × i0 × ηt / r  [N]'),
    ('行驶阻力：', 'Ff + Fw = G × f + CDA × ua² / 21.15  [N]'),
    ('旋转质量换算系数：', 'δ = 1 + (ΣIw + If × ig² × i0² × ηt) / (m × r²)'),
    ('加速度：', 'a = (Ft - Ff - Fw) / (δ × m)  [m/s²]'),
    ('加速时间（图解积分法）：', 't = ∫(1/a) du = (1/3.6) × ∫(1/a) dua  [s]'),
]:
    add_para(label)
    add_formula(formula)

doc.add_heading('2. 计算结果', level=2)
add_para('(1) 驱动力-行驶阻力平衡图如图所示。')
if os.path.exists('fig1_driving_resistance_balance.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture('fig1_driving_resistance_balance.png', width=Inches(5.5))
    add_para('图1  汽车驱动力-行驶阻力平衡图', '宋体', Pt(10), True, WD_ALIGN_PARAGRAPH.CENTER)

add_para('(2) 最高车速与最大爬坡度：')
add_result_table(['项目', '结果'], [
    ['最高车速 ua_max', '99.08 km/h'],
    ['发动机最大转矩 Tq_max', '174.97 N·m @ 2044 r/min'],
    ['发动机最大功率 Pe_max', '61.73 kW @ 3864 r/min'],
    ['最大爬坡度 i', '35.02% (19.30°)'],
    ['最大爬坡度对应车速', '10.01 km/h'],
    ['最大爬坡度时附着率 Cφ2', '0.4348'],
])

add_para('(3) 加速度倒数曲线与加速时间：')
if os.path.exists('fig2_acceleration_curves.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture('fig2_acceleration_curves.png', width=Inches(5.8))
    add_para('图2  各档加速度曲线与2档加速度倒数曲线', '宋体', Pt(10), True, WD_ALIGN_PARAGRAPH.CENTER)

add_result_table(['项目', '时间'], [
    ['2档加速段时间', '6.34 s'], ['换挡时间', '0.80 s'],
    ['3档加速段时间', '9.15 s'], ['2档起步加速至70km/h总时间', '16.29 s'],
])

# 题目一代码
doc.add_heading('3. 程序代码', level=2)
add_para('文件名：part1_dynamic_performance.py    运行：conda run -n modeling python part1_dynamic_performance.py')
with open('part1_dynamic_performance.py', 'r', encoding='utf-8') as f:
    add_code_block(f.read())

add_page_break()

# ============================================================
# 题目二
# ============================================================
add_problem_cover('二', '习题2.7')
add_para('2.7 题目：习题图2-1是《课程设计一》中的轻型货车汽油发动机的负荷特性与万有特性。货车的参数与《课程设计一》相同。负荷特性曲线的拟合公式为：')
add_formula('b = B0 + B1·Pe + B2·Pe² + B3·Pe³ + B4·Pe⁴')
add_para('式中，b为燃油消耗率[g/(kW·h)]；Pe为发动机净功率(kW)。')
add_para('拟合公式中的系数如表所示。')
add_para('（注：本题最好上机计算，可以选择Matlab、Python、Excel、VB、VC等软件进行计算和画图。）')

add_result_table(['n (r/min)', 'B0', 'B1', 'B2', 'B3', 'B4'], [
    ['815', '1326.8', '-416.46', '72.379', '-5.8629', '0.17768'],
    ['1207', '1354.7', '-303.98', '36.657', '-2.0553', '0.043072'],
    ['1614', '1284.4', '-189.75', '14.524', '-0.51184', '0.0068164'],
    ['2012', '1122.9', '-121.59', '7.0035', '-0.18517', '0.0018555'],
    ['2603', '1141.0', '-98.893', '4.4763', '-0.091077', '0.00068906'],
    ['3006', '1051.2', '-73.714', '2.8593', '-0.05138', '0.00035032'],
    ['3403', '1233.9', '-84.478', '2.9788', '-0.047449', '0.00028230'],
    ['3804', '1129.7', '-45.291', '0.71113', '-0.00075215', '-0.000038568'],
])

doc.add_heading('解题过程', level=1)
doc.add_heading('1. 计算公式', level=2)
for label, formula in [
    ('发动机功率：', 'Pe = Tq × n / 9550  [kW]'),
    ('阻力功率：', 'Pf = (Ff + Fw) × ua / (3600 × ηt)  [kW]'),
    ('等速百公里油耗：', 'Qs = Pe × b / (1.02 × ua × γ)  [L/100km]'),
    ('瞬时燃油消耗量：', 'Qt = Pe × b / (3600 × ρ)  [mL/s]'),
]:
    add_para(label)
    add_formula(formula)
add_para('式中，γ = ρg = 0.715×9.8 = 7.007 N/L（93#汽油重度）；ρ = 0.715 kg/L。')

doc.add_heading('2. 计算结果', level=2)

add_para('(1) 功率平衡图如图所示。')
if os.path.exists('fig3_power_balance.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture('fig3_power_balance.png', width=Inches(5.5))
    add_para('图3  汽车功率平衡图', '宋体', Pt(10), True, WD_ALIGN_PARAGRAPH.CENTER)

add_para('(2) 等速百公里油耗曲线如图所示。')
if os.path.exists('fig4_constant_speed_fuel.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture('fig4_constant_speed_fuel.png', width=Inches(5.0))
    add_para('图4  最高档与次高档等速百公里油耗曲线', '宋体', Pt(10), True, WD_ALIGN_PARAGRAPH.CENTER)

add_para('(3) 等速百公里油耗特征值：')
add_result_table(['项目', '油耗'], [
    ['5档（最高档）最低油耗', '11.99 L/100km @ 30 km/h'],
    ['5档 60km/h等速油耗', '13.84 L/100km'],
    ['5档 80km/h等速油耗', '17.96 L/100km'],
    ['4档（次高档）最低油耗', '13.35 L/100km @ 30 km/h'],
    ['4档 60km/h等速油耗', '15.04 L/100km'],
    ['4档 80km/h等速油耗', '19.22 L/100km'],
])

add_para('(4) 六工况循环百公里油耗：')
add_result_table(['工况', '油耗'], [
    ['工况I（等速25km/h，7.2s）', '6.02 mL'],
    ['工况II（25→40km/h加速，16.7s）', '24.26 mL'],
    ['工况III（等速40km/h，22.5s）', '31.07 mL'],
    ['工况IV（40→50km/h加速，14.0s）', '23.24 mL'],
    ['工况V（等速50km/h，18.0s）', '32.25 mL'],
    ['工况VI（50→25km/h减速，19.3s）', '24.87 mL'],
    ['总油耗 / 总行程', '141.70 mL / 1075 m'],
    ['六工况百公里油耗', '13.18 L/100km'],
])

doc.add_heading('3. 程序代码', level=2)
add_para('文件名：part2_fuel_economy.py    运行：conda run -n modeling python part2_fuel_economy.py')
with open('part2_fuel_economy.py', 'r', encoding='utf-8') as f:
    add_code_block(f.read())

add_page_break()

# ============================================================
# 题目三
# ============================================================
add_problem_cover('三', '习题4.3')
add_para('4.3 一中型货车装有前后制动器分开的双管路制动系，其有关参数如下：')
add_result_table(['载荷', '质量(kg)', '质心高hg/m', '轴距L/m', '质心至前轴a/m', '制动力分配系数β'], [
    ['空载', '4080', '0.845', '3.950', '2.100', '0.38'],
    ['满载', '9290', '1.170', '3.950', '2.950', '0.38'],
])
add_para('1）计算并绘制利用附着系数曲线和制动效率曲线。')
add_para('2）求行驶车速Ua＝30km/h，在φ＝0.80路面上车轮不抱死的制动距离。计算时取制动系反应时间τ\'＝0.02s，制动减速度上升时间τ\'\'＝0.02s。')
add_para('3）求制动系前部管路损坏时汽车的制动距离s，制动系后部管路损坏时汽车的制动距离s\'。')

doc.add_heading('解题过程', level=1)
doc.add_heading('1. 计算公式', level=2)
for label, formulas in [
    ('利用附着系数：', ['φf = β × z × L / (b + z × hg)', 'φr = (1-β) × z × L / (a - z × hg)']),
    ('制动效率：', ['Ef = z / φf', 'Er = z / φr']),
    ('同步附着系数：', ['φ0 = (L × β - b) / hg']),
    ('制动距离：', ['s = (τ1 + τ2/2) × ua0 / 3.6 + ua0² / (25.92 × a_bmax)']),
    ('最大制动减速度：', ['a_bmax = E × φ × g  [m/s²]']),
]:
    add_para(label)
    for f in formulas:
        add_formula(f)
add_para('式中，z为制动强度，b = L-a为质心至后轴距离，τ1=0.02s，τ2=0.02s。')

doc.add_heading('2. 计算结果', level=2)

add_para('(1) 利用附着系数曲线和制动效率曲线如图所示。')
if os.path.exists('fig5_adhesion_efficiency.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture('fig5_adhesion_efficiency.png', width=Inches(5.8))
    add_para('图5  利用附着系数曲线与制动效率曲线', '宋体', Pt(10), True, WD_ALIGN_PARAGRAPH.CENTER)

add_para('(2) 同步附着系数：')
add_result_table(['载荷状态', '同步附着系数 φ0'], [['空载', '-0.413'], ['满载', '0.428']])

add_para('(3) 制动距离（Ua=30km/h, φ=0.80）：')
add_result_table(['工况', '制动效率 E', '最大减速度', '制动距离 s'], [
    ['满载正常制动', '0.8715', '6.83 m/s²', '5.33 m'],
    ['空载正常制动', '0.6720', '5.27 m/s²', '6.84 m'],
    ['前部管路损坏（仅后轮）', '0.6038', '4.73 m/s²', '7.59 m'],
    ['后部管路损坏（仅前轮）', '0.3318', '2.60 m/s²', '13.60 m'],
])

doc.add_heading('3. 程序代码', level=2)
add_para('文件名：part3_braking_performance.py    运行：conda run -n modeling python part3_braking_performance.py')
with open('part3_braking_performance.py', 'r', encoding='utf-8') as f:
    add_code_block(f.read())

# ============ 保存并修改主题 ============
output_path = '课程设计报告.docx'
buf = io.BytesIO()
doc.save(buf)
buf.seek(0)

zin = zipfile.ZipFile(buf, 'r')
zout_buf = io.BytesIO()
zout = zipfile.ZipFile(zout_buf, 'w', zipfile.ZIP_DEFLATED)

for item in zin.infolist():
    data = zin.read(item.filename)
    if item.filename == 'word/theme/theme1.xml':
        content = data.decode('utf-8')
        content = re.sub(r'<a:latin typeface="[^"]*"', '<a:latin typeface="宋体"', content)
        content = re.sub(r'<a:ea typeface="[^"]*"', '<a:ea typeface="宋体"', content)
        if '<a:ea ' not in content:
            content = content.replace('</a:majorFont>', '<a:ea typeface="宋体"/></a:majorFont>')
            content = content.replace('</a:minorFont>', '<a:ea typeface="宋体"/></a:minorFont>')
        data = content.encode('utf-8')
    zout.writestr(item, data)

zin.close()
zout.close()

with open(output_path, 'wb') as f:
    f.write(zout_buf.getvalue())

print(f'报告已保存：{output_path}')
